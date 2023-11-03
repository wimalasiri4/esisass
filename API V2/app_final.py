from flask import Flask, request, render_template, send_file
from urllib.request import urlopen
from bs4 import BeautifulSoup
import re
from docx import Document
import requests
from PIL import Image
from docx.shared import Inches
import io
from io import BytesIO
from docx2pdf import convert
from flask_cors import CORS

import nltk
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.probability import FreqDist
from nltk.tokenize.treebank import TreebankWordDetokenizer

import pythoncom  # Import pythoncom for COM initialization and cleanup

app = Flask(__name__)
CORS(app)

@app.route("/download", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        search_str = request.form.get("animal_name")

        # # Check if the search_str is None or empty
        if not search_str:
            return "Please enter an animal name."

        # # Create a Word document
        # document = Document()

        # Set the predefined animal name
        # search_str = "frog"

        # Initialize the COM library
        pythoncom.CoInitialize()

        # Create a Word document
        document = Document()

        def generate_summary(text, num_sentences=2):
            sentences = sent_tokenize(text)
            words = word_tokenize(text)
            stop_words = set(stopwords.words('english'))
            filtered_words = [word for word in words if word.lower() not in stop_words]
            freq_dist = FreqDist(filtered_words)

            sentence_scores = {}
            for sentence in sentences:
                for word in word_tokenize(sentence.lower()):
                    if word in freq_dist:
                        if sentence not in sentence_scores:
                            sentence_scores[sentence] = freq_dist[word]
                        else:
                            sentence_scores[sentence] += freq_dist[word]

            summary_sentences = sorted(sentence_scores, key=sentence_scores.get, reverse=True)[:num_sentences]
            summary = TreebankWordDetokenizer().detokenize(summary_sentences)
            return summary

        def find_in_wiki(animal):
            # Specify url of the web page
            document.add_heading('Wikipedia', level=1)
            source = urlopen('https://en.wikipedia.org/wiki/'+animal).read()
            
            # Make a soup 
            soup = BeautifulSoup(source,'lxml')

            images = soup.select("div img")
            imagesUrls=[]
            for img in images:
                image_url = img['src']
                try:
                    img_size=int(img['width'])
                except:
                    pass
                if image_url[-3:].lower()=='jpg' and img_size>100:
                    image_url='http:'+image_url
                    imagesUrls.append(image_url)


            for table in soup.find_all("table"):
                table.extract()

            # Extract the plain text content from paragraphs
            paras = []
            for paragraph in soup.find_all('p'):
                test_text = str(paragraph.text)
                #print(type(test_text))
                summary = generate_summary(test_text)
                paras.append(summary)

            # Extract text from paragraph headers
            heads = []
            for head in soup.find_all('span', attrs={'mw-page-title-main'}):
                heads.append(str(head.text))
            # Interleave paragraphs & headers
            text = [val for pair in zip(paras, heads) for val in pair]
            text = '\n'.join(text)

            # Drop footnote superscripts in brackets
            text = re.sub(r"\[.*?\]+", '', text)

            # Replace '\n' (a new line) with '' and end the string at $1000.
            text = text[:-11]

            document.add_heading(heads[0], level=1)
            # for i,image in enumerate(imagesUrls[:4]):
            #     img_data = requests.get(image).content 
            #     flname = 'wiki'+str(i)+'.jpg'
            #     with open(flname, 'wb') as handler: 
            #         handler.write(img_data) 

            #     image = Image.open(flname)

            #     # Convert and save the image as JPEG
            #     image.save(flname, 'JPEG')

                
            #     document.add_picture(flname)
            
            for i, image in enumerate(imagesUrls[:4]):
                try:
                    img_data = requests.get(image).content
                    img = Image.open(io.BytesIO(img_data))
                    img.verify()  # Check if it's a valid image
                    flname = 'wiki' + str(i) + '.jpg'
                    with open(flname, 'wb') as handler:
                        handler.write(img_data)

                    image = Image.open(flname)

            # Convert and save the image as JPEG
                    image.save(flname, 'JPEG')

                    document.add_picture(flname)
                except (IOError, OSError, Image.DecompressionBombError):
            # Handle invalid image files (skip them)
                    pass
            

            num_of_paras=0
            for para in paras:
                document.add_paragraph(para)
                num_of_paras+=1
                if num_of_paras==8:
                    break


        def find_in_dilmah(search_string):
            # Specify the URL of the website to scrape
            document.add_heading('DilmahConservation', level=1)
            url = "https://www.dilmahconservation.org/about-animals"

            # Send a GET request to the website
            response = requests.get(url)
            
            # Parse the HTML content using BeautifulSoup
            soup = BeautifulSoup(response.content, "html.parser")
            # print(soup)
            # Find all article elements
            content = soup.find_all("div",attrs={'class','content-inner-wrapper'})
            

            # # Iterate over the articles
            for i,div in enumerate(content):
                
                title=div.find("h3",attrs={'class','animal-profile'})
                if search_string.lower() in title.text.lower():
                    a=div.find('a')
                    link=a.attrs['href']

                    response = requests.get(link)
                    soupCont = BeautifulSoup(response.content, "html.parser")
                    images = soupCont.select("div img")
                    imagesUrls=[]
                    for img in images:
                        # print(img)
                        image_url = img['src']
                        im_name=img['alt']
                        if im_name==title.text.strip():
                    
                            imagesUrls.append(image_url)

                    contentinner = soupCont.find_all("div",attrs={'class','content-inner-wrapper'})
                
                    animalconts=contentinner[0].find_all('p')

                    document.add_heading(title.text, level=2)

                    # Get image data and save in doc
                    img_data = requests.get(imagesUrls[0]).content 
                    flname = 'dilmah'+str(i)+'.jpg'
                    with open(flname, 'wb') as handler: 
                        handler.write(img_data) 
                    handler.close()

                    image = Image.open(flname)
                    # Convert and save the image as JPEG
                    image.save(flname, 'JPEG')
                    image.close()
                    
                    document.add_picture(flname)

                    # Add the content as a paragraph in the Word document
                    for animalcont in animalconts:
                        document.add_paragraph(animalcont.text)

        def find_in_animalia(search_string):
           # Specify the URL of the website to scrape
            document.add_heading('Animalia.bio', level=1)
            url = "https://animalia.bio/elastic-search?search="+search_string

            # Send a GET request to the website
            response = requests.get(url)
            
            # Parse the HTML content using BeautifulSoup
            soup = BeautifulSoup(response.content, "html.parser")
            # print(soup)
            # Find all article elements
            content = soup.find_all("div",attrs={'class','collection-name-text'})
            
            i=0
            for div in content:
                
                title = div.text
                modLink = title.lower().replace(' ','-').replace("'", "").strip()
                link = "https://animalia.bio/"+modLink
                response = requests.get(link)
                soupCont = BeautifulSoup(response.content, "html.parser")
                images = soupCont.select("div img")
                imagesUrls=[]
                for img in images:
                    # print(img)
                    image_url = img['src']
                    im_name=img['alt']
                    if im_name==title.strip():
                        imagesUrls.append(image_url)

                contentinner = soupCont.find_all("div",attrs={'class','s-char-text'})
                try:
                    animalcont=contentinner[0].find('p')
                    try:
                        animalcont.a.unwrap()
                    except:
                        pass 
                    para = animalcont.get_text()
                    document.add_heading(title, level=2)

                    # Get image data and save in doc
                    img_data = requests.get(imagesUrls[0]).content 
                    flname = 'anomalia'+str(i)+'.jpg'
                    with open(flname, 'wb') as handler: 
                        handler.write(img_data) 
                    handler.close()
                    image = Image.open(flname)
                    # Convert and save the image as JPEG
                    image.save(flname, 'JPEG')
                    image.close()
                    document.add_picture(flname, width=Inches(6), height=Inches(4))

                    # Add the content as a paragraph in the Word document
                    document.add_paragraph(para)
                    i+=1
                    if i==4:
                        break
                except:
                    pass

        # Call functions to scrape data
        find_in_wiki(search_str)
        find_in_dilmah(search_str)
        find_in_animalia(search_str)

        # Save the Word document in memory
        # doc_name = search_str +"_scraped_articles.docx"
        doc_name = search_str +"_scraped_articles.docx"
        

        document.save(doc_name)

        # Convert the Word document to PDF
        pdf_buffer = "D:/Freelancing/SLIIT FYP/API V2/Scraped_Files"+search_str +"_scraped_articles.pdf"
        convert(doc_name, pdf_buffer)

        # Uninitialize the COM library when done
        pythoncom.CoUninitialize()

        # Return the PDF for download
        return send_file(
            pdf_buffer,
            download_name=f"{search_str}_scraped_articles.pdf",
            mimetype="application/pdf",
            as_attachment=True,
        )


if __name__ == "__main__":
    app.run(debug=True)
